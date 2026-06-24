"""
word_exporter.py — 导出 .docx 论文（python-docx）

将分析结果、图表、图注、Methods 等组合成一个 Word 文档。
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agent.state import AgentState


def export_word(state: AgentState) -> dict:
    """
    生成 Word 文档并保存到 outputs/exports/。

    Args:
        state: AgentState

    Returns:
        {"docx_path": str}
    """
    doc = Document()

    # 标题
    title = doc.add_heading("Spatial Transcriptomics Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元数据
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Data type: {state.get('data_type', 'unknown')}")
    doc.add_paragraph("")

    # 分析步"e 一一展现
    step_results = state.get("step_results", {})
    explanations = state.get("explanations", {})
    figures = state.get("figures", {})

    for step in ["qc", "preprocess", "dimred", "cluster", "spatial", "marker"]:
        if step in step_results:
            doc.add_heading(f"{step.upper()}", level=1)
            result = step_results[step]
            doc.add_paragraph(result.get("summary", ""))

            if step in explanations:
                doc.add_paragraph(explanations[step])

            if step in figures and os.path.exists(figures[step]):
                doc.add_picture(figures[step], width=Inches(5))
                doc.add_paragraph("")

    # Methods
    skill_outputs = state.get("skill_outputs", {})
    nature_out = skill_outputs.get("nature_publish", {})
    if nature_out and nature_out.get("methods"):
        doc.add_heading("Methods", level=1)
        doc.add_paragraph(nature_out["methods"])

    # Captions
    if nature_out and nature_out.get("captions"):
        doc.add_heading("Figure Legends", level=1)
        doc.add_paragraph(nature_out["captions"])

    # Save
    os.makedirs("outputs/exports", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = f"outputs/exports/report_{timestamp}.docx"
    doc.save(docx_path)

    return {"docx_path": docx_path}